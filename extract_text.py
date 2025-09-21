# utils/extract_text.py
import os
from io import BytesIO
from typing import Union

# try multiple libraries
try:
    import fitz   # PyMuPDF
except Exception:
    fitz = None

try:
    import docx
except Exception:
    docx = None

def _read_bytes(obj):
    """
    Return bytes from a path or a file-like object.
    """
    if isinstance(obj, (str, os.PathLike)):
        with open(obj, "rb") as f:
            return f.read()
    else:
        # file-like from Streamlit upload: has .read()
        return obj.read()

def extract_text(source: Union[str, bytes, object]) -> str:
    """
    Extract text from a file path or an uploaded file-like object (Streamlit).
    Supports: PDF, DOCX, TXT.
    Use: text = extract_text(uploaded_file)  OR extract_text("path/to/file.pdf")
    """
    name = None
    # detect filename if present
    try:
        name = getattr(source, "name", None)
    except Exception:
        name = None

    # if source is a path string
    if isinstance(source, (str, os.PathLike)):
        _, ext = os.path.splitext(str(source))
        ext = ext.lower()
        if ext == ".pdf":
            return _extract_pdf_from_path(str(source))
        elif ext in [".docx", ".doc"]:
            return _extract_docx_from_path(str(source))
        elif ext in [".txt", ".md"]:
            with open(source, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            return ""

    # if source is uploaded file-like (Streamlit), use .read() bytes
    data = _read_bytes(source)
    # try to infer from name
    if name:
        _, ext = os.path.splitext(name)
        ext = ext.lower()
    else:
        ext = None

    if ext == ".pdf" or (ext is None and fitz is not None and data[:4] == b"%PDF"):
        return _extract_pdf_from_bytes(data)
    if ext in [".docx", ".doc"] or (ext is None and data[:2] == b'PK'):
        return _extract_docx_from_bytes(data)
    # fallback as text
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""

# PDF helpers
def _extract_pdf_from_bytes(data: bytes) -> str:
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) is required for PDF extraction. `pip install PyMuPDF`")
    text = ""
    doc = fitz.open(stream=data, filetype="pdf")
    for page in doc:
        text += page.get_text()
    return text

def _extract_pdf_from_path(path: str) -> str:
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) is required for PDF extraction. `pip install PyMuPDF`")
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text()
    return text

# DOCX helpers
def _extract_docx_from_bytes(data: bytes) -> str:
    if docx is None:
        raise ImportError("python-docx is required for DOCX extraction. `pip install python-docx`")
    bio = BytesIO(data)
    document = docx.Document(bio)
    paragraphs = [p.text for p in document.paragraphs]
    return "\n".join(paragraphs)

def _extract_docx_from_path(path: str) -> str:
    if docx is None:
        raise ImportError("python-docx is required for DOCX extraction. `pip install python-docx`")
    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs]
    return "\n".join(paragraphs)
